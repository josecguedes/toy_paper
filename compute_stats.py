"""Descriptive statistics for the Stack Overflow 2025 Developer Survey."""

import pandas as pd
from docx import Document
from docx.shared import Pt, RGBColor, Inches, Twips
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
import warnings
warnings.filterwarnings("ignore")

CSV_PATH  = "data/results.csv"
OUT_PATH  = "descriptive_statistics.docx"

# Colour palette
C_BLUE_DARK   = "1F497D"
C_BLUE_MED    = "2E75B6"
C_BLUE_LIGHT  = "D6E4F0"
C_BLUE_ALT    = "EBF3FB"
C_GREEN_DARK  = "375623"
C_GREEN_MED   = "4CAF50"
C_GREEN_LIGHT = "E2EFDA"
C_GREEN_ALT   = "F2F9EE"
C_WHITE       = "FFFFFF"
C_GREY_HDR    = "F2F2F2"

# ---------------------------------------------------------------------------
# Load data
# ---------------------------------------------------------------------------
df = pd.read_csv(CSV_PATH, low_memory=False)
n_rows, n_cols = df.shape
print(f"Loaded {n_rows:,} rows x {n_cols} columns")

# ---------------------------------------------------------------------------
# Column groups
# ---------------------------------------------------------------------------
NUMERIC_COLS = {
    "WorkExp":             "Work experience (years at current employer)",
    "YearsCode":           "Years coding (total)",
    "CompTotal":           "Total compensation (local currency)",
    "ConvertedCompYearly": "Annual compensation (USD, converted)",
    "JobSat":              "Job satisfaction (1-10 scale)",
    "ToolCountWork":       "Number of tools used at work",
    "ToolCountPersonal":   "Number of tools used personally",
}

CATEGORICAL_COLS = {
    "MainBranch":   "Main professional branch",
    "Age":          "Age group",
    "EdLevel":      "Education level",
    "Employment":   "Employment status",
    "RemoteWork":   "Remote-work arrangement",
    "ICorPM":       "Role type (IC vs. people manager)",
    "DevType":      "Developer type",
    "OrgSize":      "Organisation size",
    "Industry":     "Industry",
    "Country":      "Country (top 15)",
    "AISelect":     "AI tool usage frequency",
    "AISent":       "Sentiment toward AI tools",
    "AIAcc":        "Perceived AI accuracy / trust",
    "AIComplex":    "AI handling of complex tasks",
    "AIAgents":     "AI agents usage",
    "AIAgentChange":"AI agents impact on workload",
    "AIThreat":     "Perceived AI job threat",
}

# ---------------------------------------------------------------------------
# Helpers — data
# ---------------------------------------------------------------------------
def to_numeric_safe(s):
    return pd.to_numeric(s, errors="coerce")

def fmt(val, decimals=2):
    if pd.isna(val):
        return "N/A"
    return f"{val:,.{decimals}f}"

def pct_str(count, total):
    return f"{count / total * 100:.1f}%"

def bar(count, total, width=12):
    filled = round(count / total * width)
    return "█" * filled + "░" * (width - filled)

# ---------------------------------------------------------------------------
# Compute numeric stats
# ---------------------------------------------------------------------------
numeric_stats = []
for col, label in NUMERIC_COLS.items():
    if col not in df.columns:
        continue
    s  = to_numeric_safe(df[col])
    ok = s.dropna()
    n_valid   = len(ok)
    n_missing = len(s) - n_valid
    if n_valid == 0:
        continue
    numeric_stats.append({
        "Variable":  label,
        "N (valid)": f"{n_valid:,}",
        "Missing":   f"{n_missing:,}",
        "Missing %": pct_str(n_missing, len(s)),
        "Mean":      fmt(ok.mean()),
        "Median":    fmt(ok.median()),
        "Std Dev":   fmt(ok.std()),
        "Min":       fmt(ok.min()),
        "Q1":        fmt(ok.quantile(0.25)),
        "Q3":        fmt(ok.quantile(0.75)),
        "Max":       fmt(ok.max()),
    })

# ---------------------------------------------------------------------------
# Compute categorical stats
# ---------------------------------------------------------------------------
cat_tables = {}
for col, label in CATEGORICAL_COLS.items():
    if col not in df.columns:
        continue
    s  = df[col].astype(str).replace("nan", pd.NA).dropna()
    vc = s.value_counts()
    top_n = 15 if col == "Country" else 20
    vc = vc.head(top_n)
    rows = []
    for val, count in vc.items():
        rows.append({
            "Category":   val,
            "Count":      count,
            "Pct":        count / len(df) * 100,
            "Pct_str":    pct_str(count, len(df)),
            "Bar":        bar(count, len(df)),
        })
    missing = df[col].isna().sum() + (df[col].astype(str) == "nan").sum()
    cat_tables[label] = {"rows": rows, "missing": missing, "total": len(df)}

# ---------------------------------------------------------------------------
# XML / style helpers
# ---------------------------------------------------------------------------
def set_cell_bg(cell, hex_color):
    tc   = cell._tc
    tcPr = tc.get_or_add_tcPr()
    shd  = OxmlElement("w:shd")
    shd.set(qn("w:val"),   "clear")
    shd.set(qn("w:color"), "auto")
    shd.set(qn("w:fill"),  hex_color)
    tcPr.append(shd)

def set_cell_margins(cell, top=60, bottom=60, left=100, right=100):
    tc   = cell._tc
    tcPr = tc.get_or_add_tcPr()
    tcMar = OxmlElement("w:tcMar")
    for side, val in [("top", top), ("bottom", bottom), ("left", left), ("right", right)]:
        node = OxmlElement(f"w:{side}")
        node.set(qn("w:w"),    str(val))
        node.set(qn("w:type"), "dxa")
        tcMar.append(node)
    tcPr.append(tcMar)

def set_col_width(table, col_idx, width_inches):
    for row in table.rows:
        cell = row.cells[col_idx]
        tc   = cell._tc
        tcPr = tc.get_or_add_tcPr()
        tcW  = OxmlElement("w:tcW")
        tcW.set(qn("w:w"),    str(int(width_inches * 1440)))
        tcW.set(qn("w:type"), "dxa")
        tcPr.append(tcW)

def style_cell(cell, font_size=9.5, bold=False, align=WD_ALIGN_PARAGRAPH.LEFT,
               color=None):
    for para in cell.paragraphs:
        para.alignment = align
        for run in para.runs:
            run.font.size = Pt(font_size)
            run.font.name = "Calibri"
            run.bold      = bold
            if color:
                run.font.color.rgb = color

def set_row_height(row, height_inches=0.25):
    tr   = row._tr
    trPr = tr.get_or_add_trPr()
    trH  = OxmlElement("w:trHeight")
    trH.set(qn("w:val"),  str(int(height_inches * 1440)))
    trH.set(qn("w:hRule"), "atLeast")
    trPr.append(trH)

def disable_autofit(table):
    tbl  = table._tbl
    tblPr = tbl.find(qn("w:tblPr"))
    if tblPr is None:
        tblPr = OxmlElement("w:tblPr")
        tbl.insert(0, tblPr)
    tblW = OxmlElement("w:tblW")
    tblW.set(qn("w:w"),    "0")
    tblW.set(qn("w:type"), "auto")
    tblPr.append(tblW)
    layout = OxmlElement("w:tblLayout")
    layout.set(qn("w:type"), "fixed")
    tblPr.append(layout)

# ---------------------------------------------------------------------------
# Table builders
# ---------------------------------------------------------------------------
def make_numeric_table(doc, headers, rows, col_widths, header_bg, alt_bg,
                       right_cols=None):
    """Generic table with coloured header, alternating rows, fixed col widths."""
    right_cols = right_cols or []
    n_cols = len(headers)
    table  = doc.add_table(rows=1 + len(rows), cols=n_cols)
    table.style     = "Table Grid"
    table.alignment = WD_TABLE_ALIGNMENT.LEFT
    disable_autofit(table)

    # Header
    hdr = table.rows[0]
    set_row_height(hdr, 0.30)
    for i, h in enumerate(headers):
        c = hdr.cells[i]
        c.text = h
        set_cell_bg(c, header_bg)
        set_cell_margins(c, top=80, bottom=80, left=90, right=90)
        style_cell(c, font_size=9, bold=True,
                   align=WD_ALIGN_PARAGRAPH.CENTER,
                   color=RGBColor(0xFF, 0xFF, 0xFF))

    # Data rows
    for r_idx, row_data in enumerate(rows):
        r = table.rows[r_idx + 1]
        set_row_height(r, 0.25)
        is_alt = (r_idx % 2 == 1)
        for c_idx, val in enumerate(row_data):
            c = r.cells[c_idx]
            c.text = str(val)
            set_cell_margins(c, top=60, bottom=60, left=90, right=90)
            align = (WD_ALIGN_PARAGRAPH.RIGHT if c_idx in right_cols
                     else WD_ALIGN_PARAGRAPH.LEFT)
            style_cell(c, font_size=9, align=align)
            if is_alt:
                set_cell_bg(c, alt_bg)

    # Column widths
    for ci, w in enumerate(col_widths):
        set_col_width(table, ci, w)

    return table


# ---------------------------------------------------------------------------
# Build Word document
# ---------------------------------------------------------------------------
doc = Document()

# Narrow page margins so wide tables breathe
section = doc.sections[0]
section.left_margin   = Inches(0.9)
section.right_margin  = Inches(0.9)
section.top_margin    = Inches(0.9)
section.bottom_margin = Inches(0.9)

# Default body font
doc.styles["Normal"].font.name = "Calibri"
doc.styles["Normal"].font.size = Pt(10.5)

# ---- Title block ----
title_p = doc.add_paragraph()
title_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
tr = title_p.add_run("Descriptive Statistics Report")
tr.bold = True
tr.font.size = Pt(20)
tr.font.name = "Calibri"
tr.font.color.rgb = RGBColor(0x1F, 0x49, 0x7D)

sub_p = doc.add_paragraph()
sub_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
sr = sub_p.add_run("Stack Overflow Annual Developer Survey 2025")
sr.font.size   = Pt(13)
sr.font.name   = "Calibri"
sr.italic      = True
sr.font.color.rgb = RGBColor(0x40, 0x40, 0x40)

doc.add_paragraph()

# meta info box — small table
meta_tbl = doc.add_table(rows=1, cols=3)
meta_tbl.style     = "Table Grid"
meta_tbl.alignment = WD_TABLE_ALIGNMENT.CENTER
disable_autofit(meta_tbl)
meta_labels = [
    ("Total responses", f"{n_rows:,}"),
    ("Total variables", str(n_cols)),
    ("Source", "survey.stackoverflow.co"),
]
for ci, (lbl, val) in enumerate(meta_labels):
    c = meta_tbl.rows[0].cells[ci]
    set_cell_bg(c, C_BLUE_LIGHT)
    set_cell_margins(c, top=100, bottom=100, left=120, right=120)
    p = c.paragraphs[0]
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r1 = p.add_run(lbl + "\n")
    r1.bold = True; r1.font.size = Pt(8.5); r1.font.name = "Calibri"
    r1.font.color.rgb = RGBColor(0x1F, 0x49, 0x7D)
    r2 = p.add_run(val)
    r2.font.size = Pt(11); r2.font.name = "Calibri"; r2.bold = True
for ci, w in enumerate([2.2, 2.2, 3.5]):
    set_col_width(meta_tbl, ci, w)

doc.add_paragraph()

# ============================
# Section 1 — Numeric
# ============================
h1 = doc.add_heading("1.  Continuous Variables", level=1)

doc.add_paragraph(
    "Summary statistics for all quantitative variables. CompTotal is in the respondent's "
    "local currency; ConvertedCompYearly has been converted to USD by Stack Overflow."
)
doc.add_paragraph()

# 1a — Sample size & completeness
doc.add_heading("1a.  Sample Size and Missing Data", level=2)
hdrs_a  = ["Variable", "N (valid)", "Missing", "Missing %"]
rows_a  = [[s["Variable"], s["N (valid)"], s["Missing"], s["Missing %"]]
            for s in numeric_stats]
widths_a = [3.8, 1.1, 1.1, 1.0]
make_numeric_table(doc, hdrs_a, rows_a, widths_a,
                   header_bg=C_BLUE_MED, alt_bg=C_BLUE_ALT,
                   right_cols=[1, 2, 3])
doc.add_paragraph()

# 1b — Distribution
doc.add_heading("1b.  Distributional Statistics", level=2)
hdrs_b  = ["Variable", "Mean", "Median", "Std Dev", "Min", "Q1", "Q3", "Max"]
rows_b  = [[s["Variable"], s["Mean"], s["Median"], s["Std Dev"],
            s["Min"], s["Q1"], s["Q3"], s["Max"]]
            for s in numeric_stats]
widths_b = [3.1, 0.95, 0.95, 0.95, 0.95, 0.85, 0.85, 0.95]
make_numeric_table(doc, hdrs_b, rows_b, widths_b,
                   header_bg=C_BLUE_MED, alt_bg=C_BLUE_ALT,
                   right_cols=list(range(1, 8)))
doc.add_paragraph()

# ============================
# Section 2 — Categorical
# ============================
h2 = doc.add_heading("2.  Categorical Variables", level=1)
doc.add_paragraph(
    "Frequency tables for key categorical variables (up to 15-20 most common "
    "values). Percentages are over all respondents. The bar gives a visual "
    "indication of relative frequency."
)
doc.add_paragraph()

hdrs_cat = ["Category", "Count", "%", "Distribution"]
for label, info in cat_tables.items():
    doc.add_heading(label, level=2)

    missing_pct = pct_str(info["missing"], info["total"])
    note_p = doc.add_paragraph()
    note_r = note_p.add_run(
        f"Missing / not answered:  {info['missing']:,}  ({missing_pct})"
    )
    note_r.font.size  = Pt(9)
    note_r.italic     = True
    note_r.font.color.rgb = RGBColor(0x60, 0x60, 0x60)

    rows_cat = [
        [r["Category"], f"{r['Count']:,}", r["Pct_str"], r["Bar"]]
        for r in info["rows"]
    ]
    widths_cat = [3.5, 0.85, 0.75, 2.3]
    tbl = make_numeric_table(
        doc, hdrs_cat, rows_cat, widths_cat,
        header_bg=C_GREEN_DARK, alt_bg=C_GREEN_ALT,
        right_cols=[1, 2]
    )

    # Colour the bar column in a muted teal and set monospace font
    for r_idx in range(1, len(tbl.rows)):
        c    = tbl.rows[r_idx].cells[3]
        para = c.paragraphs[0]
        para.alignment = WD_ALIGN_PARAGRAPH.LEFT
        for run in para.runs:
            run.font.name  = "Courier New"
            run.font.size  = Pt(8)
            run.font.color.rgb = RGBColor(0x2E, 0x75, 0xB6)

    doc.add_paragraph()

# ---- Save ----
doc.save(OUT_PATH)
print(f"Saved -> {OUT_PATH}")
